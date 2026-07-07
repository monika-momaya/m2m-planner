"""
M2M Programme Planner â€” Streamlit Web App v2
=============================================
New in v2:
  â€¢ Logo upload â†’ displayed beside event details
  â€¢ File renamed to "M2M Programme for <<Event Name>>"
  â€¢ Word doc (.docx) download option
  â€¢ Excel download (existing)

Run locally:   streamlit run m2m_app_v2.py
Deploy:        Push to GitHub â†’ share.streamlit.io
"""

import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import io, base64
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# â”€â”€ Page config â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.set_page_config(
    page_title="M2M Programme Planner",
    page_icon="ðŸ“‹",
    layout="wide",
    initial_sidebar_state="expanded",
)

# â”€â”€ Custom CSS â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.markdown("""
<style>
    .main { background-color: #FAFAFA; }
    .stApp { font-family: 'Arial', sans-serif; }
    .title-banner {
        background: linear-gradient(135deg, #7B1B1B 0%, #A52828 100%);
        color: white; padding: 1.2rem 2rem; border-radius: 10px;
        margin-bottom: 1.5rem; box-shadow: 0 4px 12px rgba(123,27,27,0.3);
        display: flex; align-items: center; gap: 1.5rem;
    }
    .title-banner-text h1 { color: white; margin: 0; font-size: 1.8rem; }
    .title-banner-text p  { color: #F5E6C8; margin: 0.3rem 0 0 0; font-size: 0.95rem; }
    .title-banner-logo img { max-height: 70px; border-radius: 6px; }
    .section-header {
        background: #C9A84C; color: white; padding: 0.5rem 1rem;
        border-radius: 6px; font-weight: bold; margin: 1rem 0 0.5rem 0; font-size: 0.95rem;
    }
    .time-pill {
        background: #F2DADA; color: #7B1B1B; padding: 0.2rem 0.6rem;
        border-radius: 20px; font-weight: bold; font-size: 0.85rem; font-family: monospace;
    }
    .script-box {
        background: #F8F9FA; border-left: 4px solid #7B1B1B;
        padding: 0.8rem 1rem; border-radius: 0 6px 6px 0;
        margin: 0.3rem 0; font-size: 0.88rem; line-height: 1.6;
    }
    .script-box-kn {
        background: #FFF8EE; border-left: 4px solid #C9A84C;
        padding: 0.8rem 1rem; border-radius: 0 6px 6px 0;
        margin: 0.3rem 0; font-size: 0.88rem; line-height: 1.6;
    }
    .summary-card {
        background: #2C2C2C; color: #C9A84C; padding: 1rem 1.5rem;
        border-radius: 8px; text-align: center; font-weight: bold;
    }
    .event-info-panel {
        background: white; border: 1px solid #F5E6C8; border-radius: 8px;
        padding: 1rem 1.5rem; margin-bottom: 1rem;
        display: flex; align-items: center; gap: 1.5rem;
    }
    .stDownloadButton > button {
        background: #7B1B1B !important; color: white !important;
        border: none !important; border-radius: 6px !important;
        font-weight: bold !important;
    }
    .stDownloadButton > button:hover { background: #A52828 !important; }
    .dl-word > button { background: #1F4E79 !important; }
    .dl-word > button:hover { background: #2E75B6 !important; }
</style>
""", unsafe_allow_html=True)

# â”€â”€ Script library â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
SCRIPTS = [
    (["welcome address","welcome speech"],
     "We will now have the Welcome Address.\n"
     "I request [Name], [Designation], to kindly address the gathering.\n"
     "[After speech] We thank [Name] for those inspiring words.",
     "à²ˆà²— à²¸à³à²µà²¾à²—à²¤ à²­à²¾à²·à²£ à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²¹à³†à²¸à²°à³], [à²¹à³à²¦à³à²¦à³†] à²…à²µà²°à²¨à³à²¨à³ à²¸à²­à³†à²¯à²¨à³à²¨à³ à²‰à²¦à³à²¦à³‡à²¶à²¿à²¸à²¿ à²®à²¾à²¤à²¨à²¾à²¡à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ à¤¹à¤® à¤¸à¥à¤µà¤¾à¤—à¤¤ à¤­à¤¾à¤·à¤£ à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¹à¤® [à¤¨à¤¾à¤®], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¸à¥‡ à¤¸à¤­à¤¾ à¤•à¥‹ à¤¸à¤‚à¤¬à¥‹à¤§à¤¿à¤¤ à¤•à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤\n"
     "[à¤­à¤¾à¤·à¤£ à¤•à¥‡ à¤ªà¤¶à¥à¤šà¤¾à¤¤] à¤¹à¤® [à¤¨à¤¾à¤®] à¤•à¤¾ à¤‡à¤¨ à¤ªà¥à¤°à¥‡à¤°à¤£à¤¾à¤¦à¤¾à¤¯à¤• à¤¶à¤¬à¥à¤¦à¥‹à¤‚ à¤•à¥‡ à¤²à¤¿à¤ à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["keynote","inaugural address","chief guest"],
     "We are privileged to have the Keynote / Inaugural Address.\n"
     "I request [Full Name & Designation]\nto kindly deliver the Address.\n"
     "[After speech] Please join me in a round of applause.",
     "à²ˆà²— à²®à³à²–à³à²¯ à²­à²¾à²·à²£ à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²ªà³‚à²°à³à²£ à²¹à³†à²¸à²°à³ à²®à²¤à³à²¤à³ à²¹à³à²¦à³à²¦à³†] à²…à²µà²°à²¨à³à²¨à³ à²­à²¾à²·à²£ à²®à²¾à²¡à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.\n"
     "[à²­à²¾à²·à²£à²¦ à²¨à²‚à²¤à²°] à²šà²ªà³à²ªà²¾à²³à³†à²¯à³Šà²‚à²¦à²¿à²—à³† à²…à²­à²¿à²¨à²‚à²¦à²¿à²¸à³‹à²£.",
     "à¤¹à¤®à¥‡à¤‚ à¤®à¥à¤–à¥à¤¯ / à¤‰à¤¦à¥à¤˜à¤¾à¤Ÿà¤¨ à¤­à¤¾à¤·à¤£ à¤•à¤¾ à¤¸à¥Œà¤­à¤¾à¤—à¥à¤¯ à¤ªà¥à¤°à¤¾à¤ªà¥à¤¤ à¤¹à¥ˆà¥¤\n"
     "à¤¹à¤® à¤®à¥à¤–à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿, [à¤ªà¥‚à¤°à¥à¤£ à¤¨à¤¾à¤® à¤à¤µà¤‚ à¤ªà¤¦à¤¨à¤¾à¤®], à¤¸à¥‡\nà¤­à¤¾à¤·à¤£ à¤¦à¥‡à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤\n"
     "[à¤­à¤¾à¤·à¤£ à¤•à¥‡ à¤ªà¤¶à¥à¤šà¤¾à¤¤] à¤•à¥ƒà¤ªà¤¯à¤¾ à¤¤à¤¾à¤²à¤¿à¤¯à¥‹à¤‚ à¤¸à¥‡ à¤‰à¤¨à¤•à¤¾ à¤…à¤­à¤¿à¤¨à¤‚à¤¦à¤¨ à¤•à¤°à¥‡à¤‚à¥¤"),

    (["welcome","dais","dignitar"],
     "Respected dignitaries, honoured guests, and friends â€” a very warm welcome to you all.\n"
     "We now request our distinguished guests to kindly proceed to the dais and take their seats.\n"
     "[MC reads out names one by one as each dignitary is escorted to the stage]",
     "à²—à³Œà²°à²µà²¾à²¨à³à²µà²¿à²¤ à²…à²¤à²¿à²¥à²¿à²—à²³à³‡, à²—à²£à³à²¯ à²®à²¹à²¨à³€à²¯à²°à³‡ à²®à²¤à³à²¤à³ à²†à²¤à³à²®à³€à²¯ à²¬à²‚à²§à³à²—à²³à³‡ â€” à²¨à²¿à²®à²—à³† à²¹à³ƒà²¤à³à²ªà³‚à²°à³à²µà²• à²¸à³à²µà²¾à²—à²¤.\n"
     "à²¨à²¾à²µà³ à²—à²£à³à²¯ à²…à²¤à²¿à²¥à²¿à²—à²³à²¨à³à²¨à³ à²µà³‡à²¦à²¿à²•à³†à²¯à²²à³à²²à²¿ à²†à²¸à³€à²¨à²°à²¾à²—à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.\n"
     "[à²ªà³à²°à²¤à²¿à²¯à³Šà²¬à³à²¬ à²…à²¤à²¿à²¥à²¿à²¯ à²¹à³†à²¸à²°à²¨à³à²¨à³ à²“à²¦à²¿ à²…à²µà²°à²¨à³à²¨à³ à²¸à³à²µà²¾à²—à²¤à²¿à²¸à²¿]",
     "à¤¸à¤®à¥à¤®à¤¾à¤¨à¤¿à¤¤ à¤…à¤¤à¤¿à¤¥à¤¿à¤—à¤£, à¤—à¤£à¤®à¤¾à¤¨à¥à¤¯ à¤®à¤¹à¤¾à¤¨à¥à¤­à¤¾à¤µ à¤”à¤° à¤®à¤¿à¤¤à¥à¤°à¤—à¤£ â€” à¤†à¤ª à¤¸à¤­à¥€ à¤•à¤¾ à¤¹à¤¾à¤°à¥à¤¦à¤¿à¤• à¤¸à¥à¤µà¤¾à¤—à¤¤ à¤¹à¥ˆà¥¤\n"
     "à¤¹à¤® à¤…à¤ªà¤¨à¥‡ à¤¸à¤®à¥à¤®à¤¾à¤¨à¤¿à¤¤ à¤…à¤¤à¤¿à¤¥à¤¿à¤¯à¥‹à¤‚ à¤¸à¥‡ à¤®à¤‚à¤š à¤ªà¤° à¤ªà¤§à¤¾à¤° à¤•à¤° à¤†à¤¸à¤¨ à¤—à¥à¤°à¤¹à¤£ à¤•à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤\n"
     "[MC à¤ªà¥à¤°à¤¤à¥à¤¯à¥‡à¤• à¤…à¤¤à¤¿à¤¥à¤¿ à¤•à¤¾ à¤¨à¤¾à¤® à¤ªà¥à¤•à¤¾à¤°à¤¤à¥‡ à¤¹à¥à¤ à¤‰à¤¨à¥à¤¹à¥‡à¤‚ à¤®à¤‚à¤š à¤¤à¤• à¤²à¥‡ à¤œà¤¾à¤à¤‚]"),

    (["naada","nadageethe","state anthem"],
     "We shall now commence with the Naada Geethe â€” the State Anthem of Karnataka.\n"
     "I request all those present to please rise.\n[Naada Geethe plays]\nThank you. Please be seated.",
     "à²¨à²¾à²µà³ à²•à²°à³à²¨à²¾à²Ÿà²• à²¨à²¾à²¡à²—à³€à²¤à³†à²¯à³Šà²‚à²¦à²¿à²—à³† à²•à²¾à²°à³à²¯à²•à³à²°à²®à²µà²¨à³à²¨à³ à²†à²°à²‚à²­à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.\n"
     "à²Žà²²à³à²²à²°à³‚ à²¦à²¯à²µà²¿à²Ÿà³à²Ÿà³ à²Žà²¦à³à²¦à³ à²¨à²¿à²²à³à²²à²¬à³‡à²•à³†à²‚à²¦à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²¨à³†.\n"
     "[à²¨à²¾à²¡à²—à³€à²¤à³† à²¹à²¾à²¡à²²à²¾à²—à³à²µà³à²¦à³]\nà²§à²¨à³à²¯à²µà²¾à²¦à²—à²³à³. à²¦à²¯à²µà²¿à²Ÿà³à²Ÿà³ à²•à³à²³à²¿à²¤à³à²•à³Šà²³à³à²³à²¿.",
     "à¤¹à¤® à¤…à¤¬ à¤¨à¤¾à¤¡à¤—à¥€à¤¤à¥‡ â€” à¤•à¤°à¥à¤¨à¤¾à¤Ÿà¤• à¤•à¥‡ à¤°à¤¾à¤œà¥à¤¯ à¤—à¥€à¤¤ à¤•à¥‡ à¤¸à¤¾à¤¥ à¤•à¤¾à¤°à¥à¤¯à¤•à¥à¤°à¤® à¤†à¤°à¤‚à¤­ à¤•à¤°à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤‰à¤ªà¤¸à¥à¤¥à¤¿à¤¤ à¤¸à¤­à¥€ à¤¸à¤¦à¤¸à¥à¤¯à¥‹à¤‚ à¤¸à¥‡ à¤–à¤¡à¤¼à¥‡ à¤¹à¥‹à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤¹à¥ˆà¥¤\n"
     "[à¤¨à¤¾à¤¡à¤—à¥€à¤¤à¥‡ à¤¬à¤œà¤¾à¤¯à¤¾ à¤œà¤¾à¤à¤—à¤¾]\nà¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦à¥¤ à¤•à¥ƒà¤ªà¤¯à¤¾ à¤¬à¥ˆà¤  à¤œà¤¾à¤à¤‚à¥¤"),

    (["national anthem","jana gana","jai hind"],
     "We will now conclude with the National Anthem.\nI request everyone to please rise.\n"
     "[National Anthem plays]\nJai Hind! Jai Karnataka!\n"
     "Thank you all. The programme now stands concluded.",
     "à²°à²¾à²·à³à²Ÿà³à²°à²—à³€à²¤à³†à²¯à³Šà²‚à²¦à²¿à²—à³† à²•à²¾à²°à³à²¯à²•à³à²°à²®à²µà²¨à³à²¨à³ à²®à³à²•à³à²¤à²¾à²¯à²—à³Šà²³à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.\n"
     "à²Žà²²à³à²²à²°à³‚ à²¦à²¯à²µà²¿à²Ÿà³à²Ÿà³ à²Žà²¦à³à²¦à³ à²¨à²¿à²²à³à²²à²¬à³‡à²•à³†à²‚à²¦à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²¨à³†.\n"
     "[à²°à²¾à²·à³à²Ÿà³à²°à²—à³€à²¤à³†]\nà²œà³ˆ à²¹à²¿à²‚à²¦à³! à²œà³ˆ à²•à²°à³à²¨à²¾à²Ÿà²•!",
     "à¤…à¤¬ à¤¹à¤® à¤°à¤¾à¤·à¥à¤Ÿà¥à¤°à¤—à¤¾à¤¨ à¤•à¥‡ à¤¸à¤¾à¤¥ à¤•à¤¾à¤°à¥à¤¯à¤•à¥à¤°à¤® à¤•à¤¾ à¤¸à¤®à¤¾à¤ªà¤¨ à¤•à¤°à¥‡à¤‚à¤—à¥‡à¥¤\nà¤¸à¤­à¥€ à¤¸à¥‡ à¤–à¤¡à¤¼à¥‡ à¤¹à¥‹à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤¹à¥ˆà¥¤\n"
     "[à¤°à¤¾à¤·à¥à¤Ÿà¥à¤°à¤—à¤¾à¤¨]\nà¤œà¤¯ à¤¹à¤¿à¤‚à¤¦!\nà¤†à¤ª à¤¸à¤­à¥€ à¤•à¤¾ à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦à¥¤ à¤•à¤¾à¤°à¥à¤¯à¤•à¥à¤°à¤® à¤•à¤¾ à¤¸à¤®à¤¾à¤ªà¤¨ à¤¹à¥‹à¤¤à¤¾ à¤¹à¥ˆà¥¤"),

    (["lamp","lighting","deepa","inauguration of"],
     "We will now proceed to the auspicious lighting of the lamp.\n"
     "I request [Chief Guest Name & Designation] and the distinguished guests\n"
     "to kindly come forward for the lamp lighting ceremony.",
     "à²ˆà²— à²¦à³€à²ª à²ªà³à²°à²œà³à²µà²²à²¨ à²•à²¾à²°à³à²¯à²•à³à²°à²® à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²®à³à²–à³à²¯ à²…à²¤à²¿à²¥à²¿ à²¹à³†à²¸à²°à³] à²®à²¤à³à²¤à³ à²—à²£à³à²¯à²°à²¨à³à²¨à³ à²¦à³€à²ª à²¬à³†à²³à²—à²¿à²¸à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ à¤¹à¤® à¤¦à¥€à¤ª à¤ªà¥à¤°à¤œà¥à¤µà¤²à¤¨ à¤•à¥‡ à¤¶à¥à¤­ à¤•à¤¾à¤°à¥à¤¯à¤•à¥à¤°à¤® à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¹à¤® [à¤®à¥à¤–à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿ à¤•à¤¾ à¤¨à¤¾à¤® à¤à¤µà¤‚ à¤ªà¤¦] à¤¤à¤¥à¤¾ à¤—à¤£à¤®à¤¾à¤¨à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿à¤¯à¥‹à¤‚ à¤¸à¥‡\n"
     "à¤¦à¥€à¤ª à¤ªà¥à¤°à¤œà¥à¤µà¤²à¤¨ à¤¹à¥‡à¤¤à¥ à¤®à¤‚à¤š à¤ªà¤° à¤ªà¤§à¤¾à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["perspective","industry","context setting","biotech","startup"],
     "We will now hear the perspective from [Name], [Designation].\n"
     "[After speech] Thank you, [Name], for those valuable insights.",
     "à²ˆà²— [à²¹à³†à²¸à²°à³], [à²¹à³à²¦à³à²¦à³†] à²…à²µà²°à²¿à²‚à²¦ à²¦à³ƒà²·à³à²Ÿà²¿à²•à³‹à²¨ à²•à³‡à²³à²²à²¿à²¦à³à²¦à³‡à²µà³†.\n"
     "[à²­à²¾à²·à²£à²¦ à²¨à²‚à²¤à²°] à²§à²¨à³à²¯à²µà²¾à²¦à²—à²³à³, [à²¹à³†à²¸à²°à³] à²…à²µà²°à²¿à²—à³†.",
     "à¤…à¤¬ à¤¹à¤® [à¤¨à¤¾à¤®], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¸à¥‡ à¤‰à¤¨à¤•à¥‡ à¤µà¤¿à¤šà¤¾à¤° à¤¸à¥à¤¨à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "[à¤­à¤¾à¤·à¤£ à¤•à¥‡ à¤ªà¤¶à¥à¤šà¤¾à¤¤] [à¤¨à¤¾à¤®] à¤•à¤¾ à¤‡à¤¨ à¤®à¥‚à¤²à¥à¤¯à¤µà¤¾à¤¨ à¤µà¤¿à¤šà¤¾à¤°à¥‹à¤‚ à¤¹à¥‡à¤¤à¥ à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦à¥¤"),

    (["introduction","recorded message","ambassador","h.e."],
     "We are privileged to have an introduction by [Name], [Designation],\n"
     "followed by a recorded message from [Speaker Name].",
     "à²ˆà²— [à²¹à³†à²¸à²°à³] à²…à²µà²°à²¿à²‚à²¦ à²ªà²°à²¿à²šà²¯ à²®à²¤à³à²¤à³ [à²¸à²‚à²¦à³‡à²¶ à²¨à³€à²¡à²¿à²¦ à²¹à³†à²¸à²°à³] à²…à²µà²° à²¸à²‚à²¦à³‡à²¶ à²ªà³à²°à²¸à³à²¤à³à²¤à²¿à²¯à²¾à²—à²²à²¿à²¦à³†.",
     "à¤¹à¤®à¥‡à¤‚ [à¤¨à¤¾à¤®], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¦à¥à¤µà¤¾à¤°à¤¾ à¤ªà¤°à¤¿à¤šà¤¯ à¤•à¤¾ à¤¸à¥Œà¤­à¤¾à¤—à¥à¤¯ à¤ªà¥à¤°à¤¾à¤ªà¥à¤¤ à¤¹à¥ˆ,\n"
     "à¤‡à¤¸à¤•à¥‡ à¤ªà¤¶à¥à¤šà¤¾à¤¤ [à¤µà¤•à¥à¤¤à¤¾ à¤•à¤¾ à¤¨à¤¾à¤®] à¤•à¤¾ à¤°à¤¿à¤•à¥‰à¤°à¥à¤¡à¥‡à¤¡ à¤¸à¤‚à¤¦à¥‡à¤¶ à¤ªà¥à¤°à¤¸à¥à¤¤à¥à¤¤ à¤•à¤¿à¤¯à¤¾ à¤œà¤¾à¤à¤—à¤¾."),

    (["release","policy","souvenir","publication","launch"],
     "We will now proceed to the release of [Name of Publication / Policy].\n"
     "I request [Chief Guest / Name] and the distinguished guests to come forward.",
     "à²ˆà²— [à²ªà³à²°à²•à²¾à²¶à²¨ / à²¨à³€à²¤à²¿ à²¹à³†à²¸à²°à³] à²¬à²¿à²¡à³à²—à²¡à³† à²®à²¾à²¡à²²à²¾à²—à³à²µà³à²¦à³.\n"
     "[à²®à³à²–à³à²¯ à²…à²¤à²¿à²¥à²¿] à²®à²¤à³à²¤à³ à²—à²£à³à²¯à²°à²¨à³à²¨à³ à²®à³à²‚à²¦à³† à²¬à²°à²²à³ à²•à³‹à²°à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ à¤¹à¤® [à¤ªà¥à¤°à¤•à¤¾à¤¶à¤¨ / à¤¨à¥€à¤¤à¤¿ à¤•à¤¾ à¤¨à¤¾à¤®] à¤•à¥‡ à¤µà¤¿à¤®à¥‹à¤šà¤¨ à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¹à¤® [à¤®à¥à¤–à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿ / à¤¨à¤¾à¤®] à¤¤à¤¥à¤¾ à¤—à¤£à¤®à¤¾à¤¨à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿à¤¯à¥‹à¤‚ à¤¸à¥‡ à¤®à¤‚à¤š à¤ªà¤° à¤ªà¤§à¤¾à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["felicitat","honour","award","memento"],
     "We will now proceed to the felicitation of our distinguished guests.\n"
     "I request [Presenter], [Designation], to kindly felicitate [Guest Name].",
     "à²ˆà²— à²—à²£à³à²¯ à²…à²¤à²¿à²¥à²¿à²—à²³ à²¸à²¨à³à²®à²¾à²¨ à²•à²¾à²°à³à²¯à²•à³à²°à²® à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²¹à³†à²¸à²°à³] à²…à²µà²°à²¨à³à²¨à³ [à²¸à²¨à³à²®à²¾à²¨à²¿à²¸à²²à³à²ªà²¡à³à²µ à²¹à³†à²¸à²°à³] à²…à²µà²°à²¨à³à²¨à³ à²¸à²¨à³à²®à²¾à²¨à²¿à²¸à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ à¤¹à¤® à¤…à¤ªà¤¨à¥‡ à¤—à¤£à¤®à¤¾à¤¨à¥à¤¯ à¤…à¤¤à¤¿à¤¥à¤¿à¤¯à¥‹à¤‚ à¤•à¥‡ à¤¸à¤®à¥à¤®à¤¾à¤¨ à¤¸à¤®à¤¾à¤°à¥‹à¤¹ à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¹à¤® [à¤ªà¥à¤°à¤¸à¥à¤¤à¥à¤¤à¤•à¤°à¥à¤¤à¤¾], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¸à¥‡ [à¤…à¤¤à¤¿à¤¥à¤¿ à¤•à¤¾ à¤¨à¤¾à¤®] à¤•à¥‹ à¤¸à¤®à¥à¤®à¤¾à¤¨à¤¿à¤¤ à¤•à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["cultural","dance","music","performance","song"],
     "We will now be treated to a cultural performance by [Name / Group].\nPlease enjoy.",
     "à²ˆà²— [à²•à²²à²¾à²µà²¿à²¦ / à²¤à²‚à²¡] à²…à²µà²°à²¿à²‚à²¦ à²¸à²¾à²‚à²¸à³à²•à³ƒà²¤à²¿à²• à²•à²¾à²°à³à²¯à²•à³à²°à²® à²ªà³à²°à²¸à³à²¤à³à²¤à²¿à²¯à²¾à²—à²²à²¿à²¦à³†.\nà²¦à²¯à²µà²¿à²Ÿà³à²Ÿà³ à²†à²¨à²‚à²¦à²¿à²¸à²¿.",
     "à¤…à¤¬ à¤¹à¤® [à¤¨à¤¾à¤® / à¤¸à¤®à¥‚à¤¹] à¤¦à¥à¤µà¤¾à¤°à¤¾ à¤à¤• à¤¸à¤¾à¤‚à¤¸à¥à¤•à¥ƒà¤¤à¤¿à¤• à¤ªà¥à¤°à¤¸à¥à¤¤à¥à¤¤à¤¿ à¤•à¤¾ à¤†à¤¨à¤‚à¤¦ à¤²à¥‡à¤‚à¤—à¥‡à¥¤\nà¤•à¥ƒà¤ªà¤¯à¤¾ à¤†à¤¨à¤‚à¤¦ à¤‰à¤ à¤¾à¤à¤‚à¥¤"),

    (["vote of thanks","vote of thank"],
     "We will now have the Vote of Thanks.\n"
     "I request [Name], [Designation], to kindly propose the Vote of Thanks.",
     "à²ˆà²— à²µà²‚à²¦à²¨à²¾à²°à³à²ªà²£à³† à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²¹à³†à²¸à²°à³], [à²¹à³à²¦à³à²¦à³†] à²…à²µà²°à²¨à³à²¨à³ à²µà²‚à²¦à²¨à²¾à²°à³à²ªà²£à³† à²¸à²²à³à²²à²¿à²¸à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ à¤¹à¤® à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦ à¤ªà¥à¤°à¤¸à¥à¤¤à¤¾à¤µ à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¹à¤® [à¤¨à¤¾à¤®], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¸à¥‡ à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦ à¤ªà¥à¤°à¤¸à¥à¤¤à¤¾à¤µ à¤ªà¥à¤°à¤¸à¥à¤¤à¥à¤¤ à¤•à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["tea","coffee","lunch","break","networking","refreshment"],
     "We will now take a short break. The next session commences at [Time].\n"
     "Refreshments are available at [Location].",
     "à²¨à²¾à²µà³ à²ˆà²— à²µà²¿à²°à²¾à²® à²¤à³†à²—à³†à²¦à³à²•à³Šà²³à³à²³à²²à²¿à²¦à³à²¦à³‡à²µà³†.\nà²®à³à²‚à²¦à²¿à²¨ à²…à²§à²¿à²µà³‡à²¶à²¨ [à²¸à²®à²¯]à²•à³à²•à³† à²ªà³à²°à²¾à²°à²‚à²­à²µà²¾à²—à²²à²¿à²¦à³†.",
     "à¤…à¤¬ à¤¹à¤® à¤à¤• à¤¸à¤‚à¤•à¥à¤·à¤¿à¤ªà¥à¤¤ à¤µà¤¿à¤°à¤¾à¤® à¤²à¥‡à¤‚à¤—à¥‡à¥¤ à¤…à¤—à¤²à¤¾ à¤¸à¤¤à¥à¤° [à¤¸à¤®à¤¯] à¤ªà¤° à¤†à¤°à¤‚à¤­ à¤¹à¥‹à¤—à¤¾à¥¤\n"
     "à¤œà¤²à¤ªà¤¾à¤¨ [à¤¸à¥à¤¥à¤¾à¤¨] à¤ªà¤° à¤‰à¤ªà¤²à¤¬à¥à¤§ à¤¹à¥ˆà¥¤"),

    (["panel","discussion","roundtable","session"],
     "We will now move to the Panel Discussion.\n"
     "Moderator: [Name, Designation]. I request the panellists to take their seats.",
     "à²ˆà²— à²¸à²®à²¿à²¤à²¿ à²šà²°à³à²šà³† à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\nà²¨à²¿à²°à³à²µà²¾à²¹à²•à²°à³: [à²¹à³†à²¸à²°à³, à²¹à³à²¦à³à²¦à³†].",
     "à¤…à¤¬ à¤¹à¤® à¤ªà¥ˆà¤¨à¤² à¤šà¤°à¥à¤šà¤¾ à¤•à¥€ à¤“à¤° à¤¬à¤¢à¤¼à¥‡à¤‚à¤—à¥‡à¥¤\n"
     "à¤¸à¤‚à¤šà¤¾à¤²à¤•: [à¤¨à¤¾à¤®, à¤ªà¤¦à¤¨à¤¾à¤®]à¥¤ à¤¹à¤® à¤ªà¥ˆà¤¨à¤²à¤¿à¤¸à¥à¤Ÿà¥à¤¸ à¤¸à¥‡ à¤…à¤ªà¤¨à¥‡ à¤†à¤¸à¤¨ à¤—à¥à¤°à¤¹à¤£ à¤•à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤"),

    (["address by","address from","address"],
     "We will now have an address by [Name], [Designation].\n"
     "I request [Name] to kindly come forward.\n[After speech] Thank you, [Name].",
     "à²ˆà²— [à²¹à³†à²¸à²°à³], [à²¹à³à²¦à³à²¦à³†] à²…à²µà²°à²¿à²‚à²¦ à²­à²¾à²·à²£ à²¨à²¡à³†à²¯à²²à²¿à²¦à³†.\n"
     "[à²¹à³†à²¸à²°à³] à²…à²µà²°à²¨à³à²¨à³ à²®à³à²‚à²¦à³† à²¬à²°à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†.",
     "à¤…à¤¬ [à¤¨à¤¾à¤®], [à¤ªà¤¦à¤¨à¤¾à¤®] à¤¦à¥à¤µà¤¾à¤°à¤¾ à¤¸à¤‚à¤¬à¥‹à¤§à¤¨ à¤¹à¥‹à¤—à¤¾à¥¤\n"
     "à¤¹à¤® [à¤¨à¤¾à¤®] à¤¸à¥‡ à¤®à¤‚à¤š à¤ªà¤° à¤ªà¤§à¤¾à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤•à¤°à¤¤à¥‡ à¤¹à¥ˆà¤‚à¥¤\n[à¤­à¤¾à¤·à¤£ à¤•à¥‡ à¤ªà¤¶à¥à¤šà¤¾à¤¤] à¤§à¤¨à¥à¤¯à¤µà¤¾à¤¦, [à¤¨à¤¾à¤®]à¥¤"),
]

LANGUAGE_OPTIONS = {
    "English + à²•à²¨à³à²¨à²¡ (Kannada)": {"code": "kan", "label": "à²•à²¨à³à²¨à²¡", "font": "Nirmala UI"},
    "English + à¤¹à¤¿à¤¨à¥à¤¦à¥€ (Hindi)":   {"code": "hin", "label": "à¤¹à¤¿à¤¨à¥à¤¦à¥€", "font": "Nirmala UI"},
}

import re as _re

def extract_name_designation(item):
    """
    Best-effort extraction of a person's Name + full Designation from a
    programme item string. Handles common patterns used in M2M sheets:

      "... by Dr. Kiran Mazumdar Shaw, Chairperson, Vision Group on
       Biotechnology, Government of Karnataka"
      "... by Shri Darshan H.V., IAS, Managing Director, KITS"
      "Address by Hon'ble Deputy Chief Minister of Karnataka"
      "Introduction by H.E. Mr. Kimmo Lahdevirta, Ambassador of Finland"

    Returns (name_with_title, designation) â€” designation includes
    everything after the first comma. If no "by" pattern is found,
    returns (None, None) and the generic template is used instead.
    """
    if not item:
        return None, None

    # Look for "by <rest>" â€” case-insensitive, handles "followed by", "and Address by" etc.
    # We want the LAST "by" in the string (closest to the actual name)
    matches = list(_re.finditer(r'\bby\b', item, flags=_re.IGNORECASE))
    if not matches:
        return None, None

    last_by = matches[-1]
    rest = item[last_by.end():].strip()
    if not rest:
        return None, None

    # Cut off at sentence-ending punctuation if present (rare in these sheets)
    rest = rest.rstrip('.').strip()

    # Split into Name (first segment before first comma) and Designation (everything after)
    if ',' in rest:
        name_part, designation = rest.split(',', 1)
        name_part = name_part.strip()
        designation = designation.strip()
    else:
        name_part = rest.strip()
        designation = ""

    if not name_part:
        return None, None

    return name_part, designation


def fill_placeholders(template, name, designation):
    """
    Replaces [Name], [Designation], [Full Name & Designation],
    [Name / Designation], [Name, Designation] etc. with the actual
    extracted values. Also handles Kannada [à²¹à³†à²¸à²°à³]/[à²¹à³à²¦à³à²¦à³†] and
    Hindi [à¤¨à¤¾à¤®]/[à¤ªà¤¦à¤¨à¤¾à¤®] placeholder variants used in regional templates.
    Falls back gracefully if designation is empty.
    """
    if not name:
        return template

    full = f"{name}, {designation}" if designation else name

    # Order matters â€” replace combined placeholders before single ones
    replacements = [
        (r'\[Full Name & Designation\]', full),
        (r'\[Name & Designation\]', full),
        (r'\[Name, Designation\]', full),
        (r'\[Name / Designation\]', full),
        (r'\[Name\]\s*,\s*\[Designation\]', full),
        (r'\[Name\],\s*\[Designation\]', full),
        (r'\[Presenter Name\]', name),
        (r'\[Presenter\]', name),
        (r'\[Guest Name\]', name),
        (r'\[Chief Guest Name & Designation\]', full),
        (r'\[Chief Guest / Name\]', name),
        (r'\[Name\]', name),
        (r'\[Designation\]', designation if designation else ""),
        # Kannada placeholder variants
        (r'\[à²¹à³†à²¸à²°à³\],?\s*\[à²¹à³à²¦à³à²¦à³†\]', full),
        (r'\[à²ªà³‚à²°à³à²£ à²¹à³†à²¸à²°à³ à²®à²¤à³à²¤à³ à²¹à³à²¦à³à²¦à³†\]', full),
        (r'\[à²¹à³†à²¸à²°à³\]', name),
        (r'\[à²¹à³à²¦à³à²¦à³†\]', designation if designation else ""),
        # Hindi placeholder variants
        (r'\[à¤¨à¤¾à¤®\],?\s*\[à¤ªà¤¦à¤¨à¤¾à¤®\]', full),
        (r'\[à¤ªà¥‚à¤°à¥à¤£ à¤¨à¤¾à¤® à¤à¤µà¤‚ à¤ªà¤¦à¤¨à¤¾à¤®\]', full),
        (r'\[à¤¨à¤¾à¤®\s*/\s*à¤ªà¤¦à¤¨à¤¾à¤®\]', full),
        (r'\[à¤¨à¤¾à¤®\]', name),
        (r'\[à¤ªà¤¦à¤¨à¤¾à¤®\]', designation if designation else ""),
    ]
    result = template
    for pattern, value in replacements:
        result = _re.sub(pattern, value, result)
    # Clean up any leftover ", ," or trailing ", " from empty designation
    result = _re.sub(r',\s*,', ',', result)
    result = _re.sub(r',\s*\n', '\n', result)
    return result


def get_script(item, lang_code="kan"):
    """
    Returns (english_text, regional_text) for a programme item.
    lang_code: 'kan' for Kannada, 'hin' for Hindi.

    Automatically extracts the speaker's Name + full Designation from
    the programme item text (best-effort, pattern: "... by <Name>, <Designation>")
    and substitutes it into the matched template in place of [Name]/[Designation]
    placeholders â€” in BOTH English and the regional language.
    """
    t = item.lower()
    name, designation = extract_name_designation(item)

    for entry in SCRIPTS:
        keywords, eng, kan, hin = entry
        if any(k in t for k in keywords):
            regional = hin if lang_code == "hin" else kan
            eng_filled = fill_placeholders(eng, name, designation)
            # For regional language, only substitute the NAME (designation kept
            # in English since titles/orgs are usually proper nouns / English terms)
            regional_filled = fill_placeholders(regional, name, designation) if name else regional
            return eng_filled, regional_filled

    eng = f"We will now have â€” {item}.\nI request [Name / Designation] to kindly come forward.\n[MC Note: Add specific announcement text]"
    kan = f"à²ˆà²— â€” {item}.\n[à²¹à³†à²¸à²°à³] à²…à²µà²°à²¨à³à²¨à³ à²®à³à²‚à²¦à³† à²¬à²°à²²à³ à²µà²¿à²¨à²‚à²¤à²¿à²¸à³à²¤à³à²¤à³‡à²µà³†."
    hin = f"à¤…à¤¬ â€” {item}.\n[à¤¨à¤¾à¤® / à¤ªà¤¦à¤¨à¤¾à¤®] à¤¸à¥‡ à¤®à¤‚à¤š à¤ªà¤° à¤ªà¤§à¤¾à¤°à¤¨à¥‡ à¤•à¤¾ à¤…à¤¨à¥à¤°à¥‹à¤§ à¤¹à¥ˆà¥¤"
    eng_filled = fill_placeholders(eng, name, designation)
    regional = hin if lang_code == "hin" else kan
    regional_filled = fill_placeholders(regional, name, designation) if name else regional
    return eng_filled, regional_filled

def is_address(item):
    return any(k in item.lower() for k in ["address","keynote","remarks","speech","perspective","introduction"])

def get_category(item):
    t = item.lower()
    if any(k in t for k in ["address","keynote","remarks","speech","perspective","context","introduction"]): return "Address / Speech"
    if any(k in t for k in ["felicitat","release","souvenir","policy","publication"]): return "Felicitation / Release"
    if any(k in t for k in ["cultural","dance","music","performance"]): return "Cultural Programme"
    if any(k in t for k in ["naada","anthem","lamp","lighting","national","inaugur"]): return "Ceremonial"
    if any(k in t for k in ["welcome","dignitar","vote","thanks","break","tea","interval"]): return "Welcome / Housekeeping"
    return "General"

CAT_COLORS = {
    "Address / Speech":       "#D6E4F0",
    "Felicitation / Release": "#FDEBD0",
    "Cultural Programme":     "#F4D03F",
    "Ceremonial":             "#E8DAEF",
    "Welcome / Housekeeping": "#D5F5E3",
    "General":                "#FFFFFF",
}

# â”€â”€ Time helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def parse_time(t):
    for fmt in ["%I:%M %p","%I:%M%p","%H:%M","%I %p"]:
        try: return datetime.strptime(t.strip().upper(), fmt)
        except: pass
    return None

def fmt_time(dt): return dt.strftime("%I:%M %p").lstrip("0")
def fmt_slot(s,e): return f"{fmt_time(s)} â€“ {fmt_time(e)}"

# â”€â”€ Safe filename â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def safe_filename(name):
    import re
    return re.sub(r'[\\/*?:"<>|]','', name).strip() or "Programme"

# â”€â”€ Word doc export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def build_word(event_name, event_date, venue, rows, logo_bytes=None, lang_code="kan"):
    """
    Generates a clean Word doc using python-docx (no Node.js dependency â€”
    guaranteed to work on Streamlit Cloud and locally).
    Layout: logo (left, large) + title + date/venue/start line,
    clean white Timings | : | Programme Details table,
    MC Script starts on page 2.
    """
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.section import WD_SECTION
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        return None, f"python-docx not available: {e}"

    try:
        doc = DocxDoc()
        # Set default document font to Calibri
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.font.name = 'Calibri'
        # Ensure Calibri applies to complex scripts (Devanagari/Kannada fallback handled per-run)
        rpr = normal_style.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rFonts.set(qn('w:eastAsia'), 'Calibri')

        for section in doc.sections:
            section.top_margin = Cm(1.6)
            section.bottom_margin = Cm(1.6)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        def rgb(hex_str):
            h = hex_str.lstrip('#')
            return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

        def set_cell_bg(cell, hex_color):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        def no_borders_table(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def light_borders(cell):
            """
            Sets borders to 'nil' so the table has NO visible printed borders â€”
            matching Word's 'View Gridlines' mode (faint on-screen guide only,
            invisible when printed or exported to PDF). Used for the Programme table.
            """
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def light_borders_visible(cell):
            """
            Sets a thin, light-grey VISIBLE border â€” used for the MC Script
            tables so each English/regional-language box is clearly outlined.
            """
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),'single')
                el.set(qn('w:sz'),'4')
                el.set(qn('w:color'),'D9D9D9')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        # â”€â”€ Header: logo (left, large) + title beside it â”€â”€
        if logo_bytes:
            hdr_tbl = doc.add_table(rows=1, cols=2)
            logo_cell, title_cell = hdr_tbl.rows[0].cells
            logo_cell.width = Cm(5.0); title_cell.width = Cm(11.6)
            no_borders_table(logo_cell); no_borders_table(title_cell)

            lp = logo_cell.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lr = lp.add_run()
            lr.add_picture(io.BytesIO(logo_bytes), width=Cm(4.5))

            tp = title_cell.paragraphs[0]
            tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tp.paragraph_format.line_spacing = Pt(24)   # Exactly 24pt = single for 20pt font
            tp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            tp.paragraph_format.space_before = Pt(0)
            tp.paragraph_format.space_after = Pt(4)
            tr = tp.add_run(f"Minute-to-Minute Programme for {event_name or 'Event'}")
            tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = rgb("1A2B4C")
            tr.font.name = 'Calibri'
        else:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.paragraph_format.line_spacing = Pt(24)
            title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(4)
            tr = title_para.add_run(f"Minute-to-Minute Programme for {event_name or 'Event'}")
            tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = rgb("1A2B4C")
            tr.font.name = 'Calibri'

        # â”€â”€ Date / Venue / Start Time line â”€â”€
        details = []
        if event_date: details.append(f"Date: {event_date}")
        if venue:      details.append(f"Venue: {venue}")
        if rows:       details.append(f"Start Time: {rows[0]['start_str']}")
        if details:
            det_para = doc.add_paragraph("  |  ".join(details))
            det_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            det_para.paragraph_format.space_after = Pt(6)
            for run in det_para.runs:
                run.font.size = Pt(10); run.font.color.rgb = rgb("555555")
                run.font.name = 'Calibri'

        # â”€â”€ Programme table: Timings | : | Programme Details â”€â”€
        # Auto-sized columns: Timings just wide enough, ":" minimal, content gets the rest
        # A4 page (21cm) - margins (1.8*2=3.6cm) = 17.4cm content width
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        col_widths = [Cm(3.8), Cm(0.4), Cm(13.2)]  # total = 17.4cm

        hdr = table.rows[0]
        for j, (cell, text) in enumerate(zip(hdr.cells, ["Timings","","Programme Details"])):
            cell.width = col_widths[j]
            light_borders(cell)
            set_cell_bg(cell, "F5F5F5")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j!=2 else WD_ALIGN_PARAGRAPH.LEFT
            if text:
                run = p.add_run(text)
                run.bold = True; run.font.size = Pt(12); run.font.color.rgb = rgb("2C2C2C")
                run.font.name = 'Calibri'

        for row in rows:
            tr_row = table.add_row()
            values = [row['slot'], ':', row['item']]
            for j, (cell, text) in enumerate(zip(tr_row.cells, values)):
                cell.width = col_widths[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                light_borders(cell)
                set_cell_bg(cell, "FFFFFF")
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run.font.color.rgb = rgb("2C2C2C")
                elif j == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.font.color.rgb = rgb("888888")
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run.font.color.rgb = rgb("2C2C2C")

        # â”€â”€ Footer: "For internal use only | <event name>" â”€â”€
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"For internal use only  |  {event_name or 'Event'}")
        footer_run.font.size = Pt(8); footer_run.font.color.rgb = rgb("999999")
        footer_run.font.name = 'Calibri'
        footer_run.italic = True

        # â”€â”€ Page break â†’ MC Script â”€â”€
        doc.add_page_break()

        mc_heading = doc.add_paragraph()
        _mc_lang_label = "à¤¹à¤¿à¤¨à¥à¤¦à¥€" if lang_code == "hin" else "à²•à²¨à³à²¨à²¡"
        hr = mc_heading.add_run(f"MC Script â€” Bilingual (English + {_mc_lang_label})")
        hr.bold = True; hr.font.size = Pt(16); hr.font.color.rgb = rgb("1A2B4C")
        hr.font.name = 'Calibri'

        note = doc.add_paragraph()
        nr = note.add_run("Replace all text in [brackets] with actual names/designations before the event.")
        nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = rgb("888888")
        nr.font.name = 'Calibri'
        doc.add_paragraph()

        for i, row in enumerate(rows):
            eng, kan = get_script(row['item'], lang_code)

            item_para = doc.add_paragraph()
            ir = item_para.add_run(f"{i+1}.  {row['item']}")
            ir.bold = True; ir.font.size = Pt(11)
            ir.font.name = 'Calibri'
            ir.font.color.rgb = rgb("1A5276") if is_address(row['item']) else rgb("1A2B4C")
            sr = item_para.add_run(f"   â€”   {row['slot']}")
            sr.font.size = Pt(10); sr.font.color.rgb = rgb("555555")
            sr.font.name = 'Calibri'

            sc_table = doc.add_table(rows=1, cols=2)
            sc_table.autofit = False
            cells = sc_table.rows[0].cells
            cells[0].width = Cm(7.7); cells[1].width = Cm(7.7)

            light_borders_visible(cells[0]); light_borders_visible(cells[1])
            set_cell_bg(cells[0], "FFFFFF")
            set_cell_bg(cells[1], "FFFFFF")

            p_eng = cells[0].paragraphs[0]
            lbl_e = p_eng.add_run("English\n")
            lbl_e.bold = True; lbl_e.font.size = Pt(9); lbl_e.font.color.rgb = rgb("1A5276")
            lbl_e.font.name = 'Calibri'
            body_e = p_eng.add_run(eng)
            body_e.font.size = Pt(9)
            body_e.font.name = 'Calibri'

            p_kan = cells[1].paragraphs[0]
            lang_label_local = "à¤¹à¤¿à¤¨à¥à¤¦à¥€" if lang_code == "hin" else "à²•à²¨à³à²¨à²¡"
            lbl_k = p_kan.add_run(lang_label_local + "\n")
            lbl_k.bold = True; lbl_k.font.size = Pt(9); lbl_k.font.color.rgb = rgb("7B5200")
            lbl_k.font.name = 'Calibri'
            body_k = p_kan.add_run(kan)
            body_k.font.size = Pt(9)
            body_k.font.name = "Nirmala UI"

            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf, None

    except Exception as e:
        import traceback
        return None, f"{type(e).__name__}: {e}\n{traceback.format_exc()[-500:]}"


# â”€â”€ MC Document builder â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def extract_all_names(rows):
    """
    Extract every person's name from ALL programme items â€” not just "by" patterns.
    Looks for titles: Shri, Smt, Dr., Mr., Ms., Mrs., Prof., H.E., Hon'ble, Sri.
    Returns list of (name_with_title, designation, source_item).
    """
    import re as _re2
    seen = set()
    results = []
    title_pattern = _re2.compile(
        r"(?:Shri|Smt\.?|Dr\.?|Mr\.?|Ms\.?|Mrs\.?|Prof\.?|H\.E\.?|Hon['']?ble|Sri|"
        r"Padma(?:shri|bhushan|vibhushan)|Padma Shri)\s+[A-Z][A-Za-z]",
        _re2.IGNORECASE
    )
    by_pattern = _re2.compile(r"\bby\b", _re2.IGNORECASE)

    for row in rows:
        item = row.get('item', '')
        if not item: continue

        # Method 1: "by <Name>, <designation>" extraction
        by_matches = list(by_pattern.finditer(item))
        if by_matches:
            rest = item[by_matches[-1].end():].strip().rstrip('.')
            if rest:
                if ',' in rest:
                    name_part, desig = rest.split(',', 1)
                else:
                    name_part, desig = rest, ''
                name_part = name_part.strip(); desig = desig.strip()
                if name_part and name_part.lower() not in seen:
                    seen.add(name_part.lower())
                    results.append((name_part, desig, item))

        # Method 2: title-based scan (catches names not after "by")
        for m in title_pattern.finditer(item):
            # Extract from match start to next sentence break / end
            fragment = item[m.start():]
            # Stop at certain conjunctions
            stop = _re2.search(r'\b(?:and|followed by|who|will|to|for)\b', fragment, _re2.I)
            if stop:
                fragment = fragment[:stop.start()]
            fragment = fragment.strip().rstrip(',.')
            if ',' in fragment:
                name_part, desig = fragment.split(',', 1)
                name_part = name_part.strip(); desig = desig.strip()
            else:
                name_part = fragment; desig = ''
            name_lower = name_part.lower()
            # Skip if this name (or its key words) already appear in a previously extracted name or designation
            already_covered = any(
                name_lower in s or s in name_lower or
                name_lower in ' '.join([r[0]+' '+r[1] for r in results]).lower()
                for s in seen
            )
            if name_part and not already_covered:
                seen.add(name_lower)
                results.append((name_part, desig, item))

    return results


def build_mc_doc(event_name, event_date, venue, rows, logo_bytes=None, lang_code="kan"):
    """
    Full MC Document with three sections:
      Page 1  â€” Dignitaries on the Dais
      Page 2+ â€” Bilingual MC Script (English + regional)
      Final   â€” Notes (only items that have a Remarks/Speaker entry)
    """
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        return None, f"python-docx not available: {e}"

    try:
        doc = DocxDoc()
        for section in doc.sections:
            section.top_margin    = Cm(1.6)
            section.bottom_margin = Cm(1.6)
            section.left_margin   = Cm(1.8)
            section.right_margin  = Cm(1.8)

        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

        NAVY = "1A2B4C"

        def rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

        def set_cell_bg(cell, hex_color):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

        def no_borders(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def light_borders(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),'nil')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def visible_borders(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'),'single')
                el.set(qn('w:sz'),'4')
                el.set(qn('w:color'),'D9D9D9')
                tcBorders.append(el)
            tcPr.append(tcBorders)

        def add_footer(section, event_name):
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = fp.add_run(f"For internal use only  |  {event_name or 'Event'}")
            fr.font.size = Pt(8); fr.font.color.rgb = rgb("999999"); fr.italic = True

        add_footer(doc.sections[0], event_name)

        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        # PAGE 1 â€” DIGNITARIES ON THE DAIS
        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after  = Pt(8)
        hr = heading.add_run("Dignitaries on the Dais")
        hr.bold = True; hr.font.size = Pt(20); hr.font.color.rgb = rgb(NAVY)
        hr.font.name = 'Calibri'

        subh = doc.add_paragraph()
        subh.paragraph_format.space_after = Pt(12)
        sr = subh.add_run(f"{event_name or 'Event'}   |   {event_date or ''}   |   {venue or ''}")
        sr.font.size = Pt(10); sr.font.color.rgb = rgb("555555"); sr.italic = True
        sr.font.name = 'Calibri'

        all_names = extract_all_names(rows)

        if all_names:
            dais_table = doc.add_table(rows=1, cols=2)
            dais_table.autofit = False
            hdr_cells = dais_table.rows[0].cells
            hdr_cells[0].width = Cm(5.5); hdr_cells[1].width = Cm(11.9)
            for cell, text in zip(hdr_cells, ["Name", "Designation / Title"]):
                set_cell_bg(cell, "F5F5F5")
                visible_borders(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                r = p.add_run(text)
                r.bold = True; r.font.size = Pt(11); r.font.name = 'Calibri'

            for i, (name, desig, source) in enumerate(all_names):
                tr = dais_table.add_row()
                bg = "FFFFFF" if i%2==0 else "FAFAFA"
                tr.cells[0].width = Cm(5.5); tr.cells[1].width = Cm(11.9)
                for cell, text in zip(tr.cells, [name, desig]):
                    set_cell_bg(cell, bg)
                    visible_borders(cell)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    run = p.add_run(text)
                    run.font.size = Pt(11); run.font.name = 'Calibri'
        else:
            no_names = doc.add_paragraph()
            no_names.add_run("No names detected in the programme. Please add speaker names to the programme items.")

        doc.add_page_break()

        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        # PAGE 2+ â€” BILINGUAL MC SCRIPT
        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        _mc_lang_label = "à¤¹à¤¿à¤¨à¥à¤¦à¥€" if lang_code == "hin" else "à²•à²¨à³à²¨à²¡"
        mc_heading = doc.add_paragraph()
        mc_heading.paragraph_format.space_before = Pt(0)
        mc_heading.paragraph_format.space_after  = Pt(6)
        mhr = mc_heading.add_run(f"MC Script â€” Bilingual (English + {_mc_lang_label})")
        mhr.bold = True; mhr.font.size = Pt(16); mhr.font.color.rgb = rgb(NAVY)
        mhr.font.name = 'Calibri'

        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_after = Pt(12)
        nr = note_p.add_run("Replace all text in [brackets] with actual names/designations before the event.")
        nr.italic = True; nr.font.size = Pt(9); nr.font.color.rgb = rgb("888888")
        nr.font.name = 'Calibri'

        for i, row in enumerate(rows):
            eng, kan = get_script(row['item'], lang_code)

            item_para = doc.add_paragraph()
            item_para.paragraph_format.space_before = Pt(10)
            item_para.paragraph_format.space_after  = Pt(4)
            ir = item_para.add_run(f"{i+1}.  {row['item']}")
            ir.bold = True; ir.font.size = Pt(11)
            ir.font.color.rgb = rgb("1A5276") if is_address(row['item']) else rgb(NAVY)
            ir.font.name = 'Calibri'
            sr2 = item_para.add_run(f"   â€”   {row['slot']}")
            sr2.font.size = Pt(10); sr2.font.color.rgb = rgb("555555")
            sr2.font.name = 'Calibri'

            sc_table = doc.add_table(rows=1, cols=2)
            sc_table.autofit = False
            cells = sc_table.rows[0].cells
            cells[0].width = Cm(8.5); cells[1].width = Cm(8.9)

            for cell in cells:
                visible_borders(cell)
                set_cell_bg(cell, "FFFFFF")

            p_eng = cells[0].paragraphs[0]
            lbl_e = p_eng.add_run("English\n")
            lbl_e.bold = True; lbl_e.font.size = Pt(9)
            lbl_e.font.color.rgb = rgb("1A5276"); lbl_e.font.name = 'Calibri'
            body_e = p_eng.add_run(eng)
            body_e.font.size = Pt(9); body_e.font.name = 'Calibri'

            p_kan = cells[1].paragraphs[0]
            lbl_k = p_kan.add_run(f"{_mc_lang_label}\n")
            lbl_k.bold = True; lbl_k.font.size = Pt(9)
            lbl_k.font.color.rgb = rgb("7B5200"); lbl_k.font.name = 'Nirmala UI'
            body_k = p_kan.add_run(kan)
            body_k.font.size = Pt(9); body_k.font.name = 'Nirmala UI'

        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        # FINAL PAGE â€” NOTES (from Remarks/Speaker column)
        # â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
        notes_rows = [(r['item'], r.get('remarks','')) for r in rows
                      if r.get('remarks','').strip()]

        if notes_rows:
            doc.add_page_break()

            notes_h = doc.add_paragraph()
            notes_h.paragraph_format.space_before = Pt(0)
            notes_h.paragraph_format.space_after  = Pt(8)
            nhr = notes_h.add_run("Notes")
            nhr.bold = True; nhr.font.size = Pt(20)
            nhr.font.color.rgb = rgb(NAVY); nhr.font.name = 'Calibri'

            notes_sub = doc.add_paragraph()
            notes_sub.paragraph_format.space_after = Pt(12)
            nsr = notes_sub.add_run(
                "Programme-specific notes and instructions for the MC.")
            nsr.font.size = Pt(10); nsr.font.color.rgb = rgb("555555")
            nsr.italic = True; nsr.font.name = 'Calibri'

            notes_table = doc.add_table(rows=1, cols=2)
            notes_table.autofit = False
            nhdr = notes_table.rows[0].cells
            nhdr[0].width = Cm(8.0); nhdr[1].width = Cm(9.4)
            for cell, text in zip(nhdr, ["Programme Item", "Notes / Instructions"]):
                set_cell_bg(cell, "F5F5F5"); visible_borders(cell)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(text)
                run.bold = True; run.font.size = Pt(11); run.font.name = 'Calibri'

            for i, (item, note) in enumerate(notes_rows):
                tr = notes_table.add_row()
                bg = "FFFFFF" if i%2==0 else "FAFAFA"
                tr.cells[0].width = Cm(8.0); tr.cells[1].width = Cm(9.4)
                for cell, text in zip(tr.cells, [item, note]):
                    set_cell_bg(cell, bg); visible_borders(cell)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    run = p.add_run(text)
                    run.font.size = Pt(11); run.font.name = 'Calibri'

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf, None

    except Exception as e:
        import traceback
        return None, f"{type(e).__name__}: {e}\n{traceback.format_exc()[-500:]}"


# â”€â”€ Excel export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def build_excel(event_name, event_date, venue, rows, logo_bytes=None, lang_code="kan"):
    MAROON="7B1B1B"; GOLD="C9A84C"; WHITE="FFFFFF"; DARK="2C2C2C"
    CREAM="FFF8EE"; LT_GOLD="F5E6C8"; LT_MAROON="F2DADA"; GREY="F7F7F7"

    def sd(s="thin",c=GOLD): return Side(style=s,color=c)
    def bdr(s="thin"): b=sd(s); return Border(left=b,right=b,top=b,bottom=b)
    def fill(h): return PatternFill("solid",fgColor=h)
    def fnt(bold=False,color=DARK,size=10,italic=False):
        return Font(name="Arial",bold=bold,color=color,size=size,italic=italic)
    def aln(h="center",v="center",wrap=False):
        return Alignment(horizontal=h,vertical=v,wrap_text=wrap)
    def mbdr():
        m=Side(style="medium",color=MAROON)
        return Border(left=m,right=m,top=m,bottom=m)

    wb = Workbook()

    # â”€â”€ Sheet 1: Programme Planner â”€â”€
    ws1=wb.active; ws1.title="Programme Planner"
    for col,w in {"A":22,"B":52,"C":18}.items():
        ws1.column_dimensions[col].width=w

    # â”€â”€ Header: Left = event info, Right = logo â”€â”€
    # Row 1: Event Name (A1) | Logo (C1 spanning down)
    ws1.row_dimensions[1].height = 26
    ws1.row_dimensions[2].height = 18
    ws1.row_dimensions[3].height = 18
    ws1.row_dimensions[4].height = 18

    # Event name â€” left side
    c=ws1["A1"]
    c.value = f"M2M Programme for Inaugural of {event_name or 'Event'}"
    c.font=Font(name="Arial",bold=True,color=WHITE,size=14)
    c.fill=fill(MAROON); c.alignment=aln(h="left")
    ws1.merge_cells("A1:B1")
    c2=ws1["C1"]; c2.fill=fill(MAROON)

    # Event details rows
    details_rows = [
        ("Event :", event_name or "â€”"),
        ("Date  :", event_date or "â€”"),
        ("Venue :", venue or "â€”"),
        ("Start Time :", rows[0]["start_str"] if rows else "â€”"),
    ]
    for di, (lbl, val) in enumerate(details_rows):
        r = 2 + di
        ws1.row_dimensions[r].height = 18  # kept tight to match logo block
        cl = ws1.cell(row=r, column=1)
        cl.value = lbl
        cl.font = Font(name="Arial",bold=True,color=MAROON,size=9)
        cl.fill = fill(LT_GOLD); cl.alignment = aln(h="right")
        cv = ws1.cell(row=r, column=2)
        cv.value = val
        cv.font = Font(name="Arial",color=DARK,size=9)
        cv.fill = fill(WHITE); cv.border = bdr()
        cv.alignment = aln(h="left")

    # Logo â€” right side (col C), sized to match the 4-row header block exactly
    if logo_bytes:
        from openpyxl.drawing.image import Image as XLImage
        try:
            img = XLImage(io.BytesIO(logo_bytes))
            # Header block total height = row1(28)+row2-4(18*3)=82pt â‰ˆ 109px (1pt=1.333px)
            img.height = 95; img.width = 170; img.anchor = "C1"
            ws1.add_image(img)
            ws1.column_dimensions["C"].width = 24
        except Exception:
            pass

    # (event details now in header rows above)

    ws1.row_dimensions[6].height=24
    for ci,hdr in enumerate(["Time Slot","Programme Item / Activity"],1):
        c=ws1.cell(row=6,column=ci)
        c.value=hdr; c.font=Font(name="Arial",bold=True,color=WHITE,size=10)
        c.fill=fill(MAROON); c.alignment=aln(wrap=True); c.border=bdr()

    from openpyxl.styles import PatternFill as PF

    # â”€â”€ LIVE FORMULA ENGINE â”€â”€
    # Hidden helper columns: D = Duration (mins), E = Start (formula), F = End (formula)
    # Column A (visible) = Time Slot, built via formula from E & F
    # This means: if the app is unreachable, the team can edit Column D (Duration)
    # or the Start Time cell directly in Excel, and Column A recalculates automatically.
    ws1.column_dimensions["D"].width=12
    ws1.column_dimensions["E"].width=2
    ws1.column_dimensions["F"].width=2

    # Hidden start-time seed cell (Excel TIME value) â€” feeds the first formula row
    start_time_str = rows[0]['start_str'] if rows else "10:00 AM"
    seed_cell = ws1["F1"]
    seed_cell.value = f"=TIMEVALUE(\"{start_time_str}\")"
    seed_cell.number_format = "h:mm AM/PM"
    ws1.column_dimensions["F"].width = 0  # hide helper column visually (width 0)

    # Duration header
    dur_hdr = ws1.cell(row=6, column=4)
    dur_hdr.value = "Duration\n(mins)"
    dur_hdr.font=Font(name="Arial",bold=True,color=WHITE,size=10)
    dur_hdr.fill=fill(MAROON); dur_hdr.alignment=aln(wrap=True); dur_hdr.border=bdr()

    for i,row in enumerate(rows):
        r=7+i; ws1.row_dimensions[r].height=22
        item=row.get('item',''); addr=is_address(item)
        duration = row.get('duration', 0)

        # Column D: Duration (mins) â€” EDITABLE, drives the formula chain
        dcell = ws1.cell(row=r, column=4)
        dcell.value = duration
        dcell.font = fnt(color=MAROON, bold=True)
        dcell.fill = PF("solid", fgColor="FFF8EE")
        dcell.alignment = aln(h="center")
        dcell.border = bdr()

        # Column E: Start time (formula) â€” first row uses seed, others chain from previous End
        ecell = ws1.cell(row=r, column=5)
        if i == 0:
            ecell.value = "=$F$1"
        else:
            ecell.value = f"=F{r-1}"
        ecell.number_format = "h:mm AM/PM"
        ecell.font = Font(size=1, color="FFFFFF")  # invisible helper

        # Column F: End time (formula) = Start + Duration/1440
        fcell = ws1.cell(row=r, column=6)
        fcell.value = f"=E{r}+D{r}/1440"
        fcell.number_format = "h:mm AM/PM"
        fcell.font = Font(size=1, color="FFFFFF")  # invisible helper

        # Column A: Time Slot â€” LIVE FORMULA combining E and F as text
        acell = ws1.cell(row=r, column=1)
        acell.value = f'=TEXT(E{r},"h:mm AM/PM")&" \u2013 "&TEXT(F{r},"h:mm AM/PM")'
        acell.fill = PF("solid", fgColor="FFFFFF")
        acell.font = fnt(color=DARK)
        acell.alignment = aln(h="left", wrap=False)
        acell.border = bdr()

        # Column B: Programme Item
        bcell = ws1.cell(row=r, column=2)
        bcell.value = item
        bcell.fill = PF("solid", fgColor="FFFFFF")
        bcell.font = fnt(color=DARK, bold=addr)
        bcell.alignment = aln(h="left", wrap=True)
        bcell.border = bdr()

    tr=7+len(rows); ws1.row_dimensions[tr].height=22
    last_row = 7+len(rows)-1
    c=ws1[f"A{tr}"]
    c.value = f'="Programme ends at "&TEXT(F{last_row},"h:mm AM/PM")&"  (auto-calculated)"'
    c.font=Font(name="Arial",bold=True,color=GOLD,size=10)
    c.fill=fill(DARK); c.alignment=aln(h="left"); c.border=mbdr()
    ws1.merge_cells(f"A{tr}:B{tr}")
    ws1.freeze_panes="A7"

    # Note about live editing
    note_row = tr + 2
    ws1.merge_cells(f"A{note_row}:B{note_row}")
    nc = ws1[f"A{note_row}"]
    nc.value = ("\u2139\ufe0f  Edit the Duration column (D) or Start Time (cell F1) â€” "
                "Time Slots above recalculate automatically. Use this sheet as an offline "
                "backup if the web app is unavailable.")
    nc.font = fnt(italic=True, color="888888", size=8)
    nc.alignment = aln(h="left", wrap=True)
    ws1.row_dimensions[note_row].height = 28

    # â”€â”€ Sheet 2: Print View (formulas reference Programme Planner) â”€â”€
    ws2=wb.create_sheet("Print View")
    ws2.column_dimensions["A"].width=22
    ws2.column_dimensions["B"].width=52
    ws2.column_dimensions["C"].width=22
    PP = "'Programme Planner'"   # sheet reference prefix

    ws2.row_dimensions[1].height=26
    c=ws2["A1"]
    c.value=f"Minute-to-Minute Programme for Inaugural of {event_name or 'Event'}"
    c.font=Font(name="Arial",bold=True,color=WHITE,size=14)
    c.fill=fill(MAROON); c.alignment=aln(h="left")
    ws2.merge_cells("A1:B1")

    # Event details â€” pull from Planner header cells via formula
    detail_labels = ["Event :","Date  :","Venue :","Start Time :"]
    for di,lbl in enumerate(detail_labels):
        r=2+di; ws2.row_dimensions[r].height=18
        cl=ws2.cell(row=r,column=1); cl.value=lbl
        cl.font=Font(name="Arial",bold=True,color=MAROON,size=9)
        cl.fill=fill(LT_GOLD); cl.alignment=aln(h="right")
        cv=ws2.cell(row=r,column=2)
        # Reference the Planner's detail cells (rows 2-5 in Planner's header)
        cv.value=f"={PP}!B{2+di}"
        cv.font=Font(name="Arial",color=DARK,size=9)
        cv.fill=fill(WHITE); cv.border=bdr(); cv.alignment=aln(h="left")

    if logo_bytes:
        from openpyxl.drawing.image import Image as XLImage
        try:
            img2=XLImage(io.BytesIO(logo_bytes))
            img2.height=95; img2.width=170; img2.anchor="C1"
            ws2.add_image(img2)
            ws2.column_dimensions["C"].width=24
        except Exception: pass

    ws2.row_dimensions[6].height=22
    for ci,hdr in enumerate(["Time Slot","Programme Item / Activity"],1):
        c=ws2.cell(row=6,column=ci)
        c.value=hdr; c.font=Font(name="Arial",bold=True,color=WHITE,size=11)
        c.fill=fill(MAROON); c.alignment=aln(); c.border=bdr()

    # Data rows â€” formula references to Programme Planner cols A (slot) and B (item)
    for i,row in enumerate(rows):
        r=7+i; ws2.row_dimensions[r].height=22
        planner_r = 7+i   # same row offset in Planner
        for ci,formula in [
            (1, f"={PP}!A{planner_r}"),   # Time Slot (formula from Planner)
            (2, f"={PP}!B{planner_r}"),   # Programme Item
        ]:
            c=ws2.cell(row=r,column=ci)
            c.value=formula
            c.fill=PF("solid",fgColor="FFFFFF")
            c.font=fnt(color=DARK)
            c.alignment=aln(h="left",wrap=(ci==2))
            c.border=bdr()
    ws2.freeze_panes="A7"

    # â”€â”€ Sheet 3: MC Script â”€â”€
    ws3=wb.create_sheet("MC Script")
    for col,w in {"A":22,"B":30,"C":50,"D":46,"E":20}.items():
        ws3.column_dimensions[col].width=w
    # MC header: info left, logo right
    ws3.row_dimensions[1].height=26
    c=ws3["A1"]
    _lang_label_xl = "à¤¹à¤¿à¤¨à¥à¤¦à¥€" if lang_code == "hin" else "à²•à²¨à³à²¨à²¡"
    c.value=f"ðŸŽ¤  MC SCRIPT â€” Bilingual (English + {_lang_label_xl})"
    c.font=Font(name="Arial",bold=True,color=WHITE,size=13)
    c.fill=fill(MAROON); c.alignment=aln(h="left")
    ws3.merge_cells("A1:D1")

    for di,(lbl,val) in enumerate([
        ("Event :", event_name or "â€”"),("Date  :", event_date or "â€”"),
        ("Venue :", venue or "â€”"),("Note  :", "Replace [brackets] with actual names")]):
        r=2+di; ws3.row_dimensions[r].height=16
        cl=ws3.cell(row=r,column=1); cl.value=lbl
        cl.font=Font(name="Arial",bold=True,color=MAROON,size=9)
        cl.fill=fill(LT_GOLD); cl.alignment=aln(h="right")
        cv=ws3.cell(row=r,column=2); cv.value=val
        cv.font=Font(name="Arial",color=DARK,size=9)
        cv.fill=fill(WHITE); cv.border=bdr(); cv.alignment=aln(h="left")
        ws3.merge_cells(f"B{r}:D{r}")

    if logo_bytes:
        from openpyxl.drawing.image import Image as XLImage
        try:
            img3=XLImage(io.BytesIO(logo_bytes))
            img3.height=95; img3.width=170; img3.anchor="E1"
            ws3.add_image(img3)
        except Exception: pass

    ws3.row_dimensions[6].height=26
    for ci,hdr in enumerate(["Time Slot","Programme Item","ðŸ“¢ English",f"ðŸ“¢ {_lang_label_xl}"],1):
        c=ws3.cell(row=6,column=ci)
        c.value=hdr; c.font=Font(name="Arial",bold=True,color=WHITE,size=10)
        c.fill=fill(MAROON); c.alignment=aln(wrap=True); c.border=bdr()
    for i,row in enumerate(rows):
        r=7+i; ws3.row_dimensions[r].height=90
        item=row.get('item',''); cat=get_category(item)
        eng,kan=get_script(item, lang_code)
        for ci,val,bg2,fs in [
            (1,row.get('slot',''),"FFFFFF",fnt(bold=True,color=MAROON,size=9)),
            (2,item,"FFFFFF",fnt(bold=is_address(item),color=DARK)),
            (3,eng,"FFFFFF",Font(name="Arial",color="1A1A2E",size=9)),
            (4,kan,"FFFFFF",Font(name="Nirmala UI",color="1A1A2E",size=9))]:
            c=ws3.cell(row=r,column=ci)
            c.value=val; c.fill=PF("solid",fgColor=bg2); c.font=fs
            c.alignment=Alignment(horizontal="center" if ci==1 else "left",
                                  vertical="top" if ci in[3,4] else "center",wrap_text=True)
            c.border=bdr()
    ws3.freeze_panes="A7"
    ws3.page_setup.orientation="landscape"

    buf=io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# MAIN APP UI
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

# â”€â”€ Sidebar â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
with st.sidebar:
    st.markdown("### âš™ï¸ Event Details")
    event_name = st.text_input("Event Name", placeholder="e.g. BTS 2025 Inauguration")
    event_date = st.text_input("Event Date", placeholder="e.g. 15-Aug-2025")
    venue      = st.text_input("Venue", placeholder="e.g. Taj Vivanta, Bengaluru")
    start_time = st.text_input("Start Time", value="10:00 AM")

    st.markdown("---")
    st.markdown("### ðŸ—£ï¸ MC Script Language")
    lang_choice = st.radio(
        "Bilingual pairing for MC Script",
        options=list(LANGUAGE_OPTIONS.keys()),
        index=0,
        help="Choose the regional language to pair with English for this event"
    )
    lang_code = LANGUAGE_OPTIONS[lang_choice]["code"]
    lang_label = LANGUAGE_OPTIONS[lang_choice]["label"]

    st.markdown("---")
    st.markdown("### ðŸ–¼ï¸ Event Logo")
    logo_file = st.file_uploader(
        "Upload logo (PNG or JPG)",
        type=["png","jpg","jpeg"],
        help="Logo will appear in downloaded Excel and Word files"
    )
    logo_bytes = logo_file.read() if logo_file else None

    st.markdown("---")
    st.markdown("### â„¹ï¸ How to use")
    st.markdown("""
1. Fill event details above
2. Upload logo (optional)
3. Add programme items + durations
4. Download Excel or Word
""")
    st.markdown("---")
    st.markdown("**Category colours:**")
    for cat, color in CAT_COLORS.items():
        if cat != "General":
            st.markdown(
                f"<span style='background:{color};padding:2px 8px;"
                f"border-radius:4px;font-size:0.78rem;color:#2C2C2C'>{cat}</span><br>",
                unsafe_allow_html=True)

# â”€â”€ Title banner (plain â€” logo goes into downloaded files only) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.markdown("""
<div class="title-banner">
    <div class="title-banner-text">
        <h1>ðŸ“‹ M2M Programme Planner</h1>
        <p>Minute-to-Minute Schedule Â· Bilingual MC Script Â· Excel & Word Download</p>
    </div>
</div>
""", unsafe_allow_html=True)

if logo_bytes:
    st.info("âœ… Logo uploaded â€” it will appear in your downloaded Excel and Word files.")

# â”€â”€ Event info panel â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
if event_name or event_date or venue:
    info_parts = []
    if event_name: info_parts.append(f"<b>{event_name}</b>")
    if event_date: info_parts.append(f"ðŸ“… {event_date}")
    if venue:      info_parts.append(f"ðŸ“ {venue}")
    if start_time: info_parts.append(f"ðŸ• {start_time}")
    st.markdown(
        f'<div class="event-info-panel">'
        f'{"  &nbsp;|&nbsp;  ".join(info_parts)}</div>',
        unsafe_allow_html=True)

# â”€â”€ Parse start time â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
start_dt = parse_time(start_time) if start_time else parse_time("10:00 AM")
if not start_dt:
    st.error(f"âš ï¸ Can't parse '{start_time}'. Use format like '10:00 AM' or '09:30 AM'")
    st.stop()

# â”€â”€ Programme table â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.markdown('<div class="section-header">ðŸ“ Programme Items</div>',
            unsafe_allow_html=True)

DEFAULT_ITEMS = [
    {"item":"MC welcomes dignitaries onto the Dias","duration":3,"remarks":""},
    {"item":"Naada Geethe (State Anthem)","duration":4,"remarks":""},
    {"item":"Welcome address by Managing Director, Karnataka Innovation and Technology Society (KITS), Dept. of Electronics, IT, Bt and S&T, Govt. of Karnataka","duration":4,"remarks":""},
    {"item":"Context setting by Secretary to Government, Department of Electronics, IT, Bt and S&T, Govt. of Karnataka","duration":5,"remarks":""},
    {"item":"Lighting of the lamp and Inauguration of BTS 2026","duration":3,"remarks":""},
    {"item":"Inaugural Address by the Chief Guest","duration":10,"remarks":""},
    {"item":"Information Technology Industry perspective by Shri Kris Gopalakrishnan, Chairperson, Vision Group on Information Technology, Govt. of Karnataka","duration":3,"remarks":""},
    {"item":"Biotechnology Industry perspective by Dr. Kiran Mazumdar Shaw, Chairperson, Vision Group on Biotechnology, Government of Karnataka","duration":3,"remarks":""},
    {"item":"Startup perspective by Shri Prashanth Prakash, Chairperson, Vision Group on Startups, Government of Karnataka","duration":3,"remarks":""},
    {"item":"Address by representative from Software Technology Parks of India (STPI)","duration":5,"remarks":""},
    {"item":"Address by Industry Leader 1","duration":7,"remarks":""},
    {"item":"Address by Industry Leader 2","duration":7,"remarks":""},
    {"item":"Address by Industry Leader 3","duration":7,"remarks":""},
    {"item":"Address by His Excellency, <country>","duration":6,"remarks":""},
    {"item":"Address by His Excellency, <country>","duration":6,"remarks":""},
    {"item":"Release of AVGC Policy and Address by","duration":5,"remarks":""},
    {"item":"Release of Biotech Policy and Address by","duration":5,"remarks":""},
    {"item":"Address by Hon'ble Minister for Information Technology & Biotechnology and Rural Development & Panchayat Raj, Govt. of Karnataka","duration":6,"remarks":""},
    {"item":"Address by Hon'ble Deputy Chief Minister of Karnataka / <State>","duration":7,"remarks":""},
    {"item":"Words of Thanks","duration":5,"remarks":""},
]

edited_df = st.data_editor(
    pd.DataFrame(DEFAULT_ITEMS),
    column_config={
        "item":     st.column_config.TextColumn("Programme Item / Activity", width="large"),
        "duration": st.column_config.NumberColumn("Duration (mins)", min_value=1, max_value=180, step=1),
        "remarks":  st.column_config.TextColumn("Remarks / Speaker", width="medium"),
    },
    num_rows="dynamic",
    use_container_width=True,
    height=430,
    key="programme_table"
)

# â”€â”€ Calculate time slots â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
rows = []
current = start_dt
for _, r in edited_df.iterrows():
    item     = str(r.get('item','')).strip() if r.get('item') else ''
    duration = r.get('duration')
    remarks  = str(r.get('remarks','')).strip() if r.get('remarks') else ''
    if not item or not duration: continue
    try: duration = int(duration)
    except: continue
    end = current + timedelta(minutes=duration)
    rows.append({
        'item': item, 'duration': duration,
        'start_str': fmt_time(current), 'end_str': fmt_time(end),
        'slot': fmt_slot(current, end), 'remarks': remarks,
    })
    current = end

# â”€â”€ Output tabs + downloads â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
if rows:
    total_mins = sum(r['duration'] for r in rows)

    # Summary cards
    c1, c2, c3 = st.columns(3)
    for col, val, label in [
        (c1, str(len(rows)),        "Programme Items"),
        (c2, f"{total_mins//60}h {total_mins%60}m", "Total Duration"),
        (c3, rows[-1]['end_str'],   "Programme Ends"),
    ]:
        with col:
            st.markdown(
                f'<div class="summary-card">'
                f'<div style="font-size:1.6rem">{val}</div>'
                f'<div style="font-size:0.8rem;color:#F5E6C8">{label}</div>'
                f'</div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    tab1, tab2, tab3 = st.tabs(["ðŸ“‹ Schedule", "ðŸŽ¤ MC Script", "ðŸ–¨ï¸ Print View"])

    with tab1:
        st.markdown('<div class="section-header">ðŸ“‹ Minute-to-Minute Schedule</div>',
                    unsafe_allow_html=True)
        for i, row in enumerate(rows):
            cat   = get_category(row['item'])
            color = CAT_COLORS.get(cat,"#FFFFFF")
            addr  = is_address(row['item'])
            st.markdown(f"""
            <div style="display:flex;align-items:center;padding:8px 14px;
                        background:{'#F9F9F9' if i%2==0 else '#FFFFFF'};
                        border-radius:4px;margin:2px 0;
                        border:1px solid #E0E0E0;">
                <span class="time-pill" style="min-width:160px;text-align:center">{row['slot']}</span>
                <span style="flex:1;margin-left:14px;font-weight:{'bold' if addr else 'normal'};
                             color:{'#1A5276' if addr else '#2C2C2C'}">{row['item']}</span>
            </div>""", unsafe_allow_html=True)

    with tab2:
        st.markdown(f'<div class="section-header">ðŸŽ¤ Bilingual MC Script (English + {lang_label})</div>',
                    unsafe_allow_html=True)
        st.info("ðŸ’¡ Replace **[bracket]** text with actual names/designations before the event")
        for i, row in enumerate(rows):
            eng, kan = get_script(row['item'], lang_code)
            with st.expander(f"**{i+1}. {row['item']}** â€” {row['slot']}", expanded=(i<2)):
                ce, ck = st.columns(2)
                with ce:
                    st.markdown("**ðŸ“¢ English**")
                    st.markdown(f'<div class="script-box">{eng.replace(chr(10),"<br>")}</div>',
                                unsafe_allow_html=True)
                with ck:
                    st.markdown(f"**ðŸ“¢ {lang_label}**")
                    st.markdown(f'<div class="script-box-kn">{kan.replace(chr(10),"<br>")}</div>',
                                unsafe_allow_html=True)

    with tab3:
        st.markdown('<div class="section-header">ðŸ–¨ï¸ Print View</div>', unsafe_allow_html=True)
        if event_name:
            st.markdown(f"### M2M Programme for {event_name}")
        info = []
        if event_date: info.append(f"ðŸ“… {event_date}")
        if venue:      info.append(f"ðŸ“ {venue}")
        if info: st.markdown("  |  ".join(info))
        st.divider()
        st.dataframe(
            pd.DataFrame([{
                "Time Slot": r['slot'],
                "Programme Item": r['item'],
            } for r in rows]),
            use_container_width=True, hide_index=True, height=460)

    # â”€â”€ Downloads â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    st.divider()
    st.markdown("### ðŸ“¥ Download")

    fname_base = f"Minute-to-Minute Programme for {safe_filename(event_name)}" if event_name else "Minute-to-Minute Programme"

    dcol1, dcol2, dcol3 = st.columns(3)

    with dcol1:
        excel_buf = build_excel(event_name, event_date, venue, rows, logo_bytes, lang_code)
        st.download_button(
            label="â¬‡ï¸ Download Excel (.xlsx)",
            data=excel_buf,
            file_name=f"{fname_base}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        st.caption("3 sheets: Programme Planner Â· Print View Â· MC Script")

    with dcol2:
        word_buf, word_err = build_word(event_name, event_date, venue, rows, logo_bytes, lang_code)
        if word_buf:
            st.download_button(
                label="â¬‡ï¸ Download M2M (.docx)",
                data=word_buf,
                file_name=f"{fname_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.caption("Programme table only â€” clean M2M for sharing")
        else:
            st.error(f"Word export failed: {word_err}")

    with dcol3:
        mc_buf, mc_err = build_mc_doc(event_name, event_date, venue, rows, logo_bytes, lang_code)
        if mc_buf:
            st.download_button(
                label="â¬‡ï¸ Download MC Document (.docx)",
                data=mc_buf,
                file_name=f"{fname_base} â€” MC Copy.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.caption("Dignitaries on the Dais Â· Bilingual MC Script Â· Notes")
        else:
            st.error(f"MC Document failed: {mc_err}")

else:
    st.info("ðŸ‘† Add programme items and durations above to generate your schedule.")

# â”€â”€ Footer â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
st.divider()
st.markdown("""
<div style="text-align:center;color:#888;font-size:0.8rem">
    M2M Programme Planner Â· Built for Event Management Teams
</div>""", unsafe_allow_html=True)
